"""① 契約書（基本契約書）の処理"""
import os
import re
import shutil
from docx import Document

from .common import clean_formatting, extract_entity_info


SEAL_CLAUSE = '本契約の成立を証するため、本書２通を作成し、甲乙署名又は記名捺印の上、各１通を保有するものとする。'
PARTNER_PAGE_KEYWORDS = ['カテゴリー及びパートナー', 'カテゴリー・パートナー']
APPENDIX2_KEYWORDS = ['別紙2', '別紙２']
APPENDIX3_KEYWORDS = ['別紙3', '別紙３']

# 別紙2の最新様式を識別するためのキーワード（最新版に含まれる文言）
APPENDIX2_LATEST_KEYWORDS = [
    '愛知・名古屋2026',  # 最新の大会名
    '2026アジア・アジアパラ競技大会',
]
# 旧様式を示すキーワード（これが含まれていたら古い）
APPENDIX2_OLD_KEYWORDS = [
    '第20回アジア競技大会',
    '2026年アジア競技大会',
]


def process_contract(filepath, output_dir, company_name, approval_type,
                     appendix2_choice, appendix2_dir):
    result = {'output_name': '', 'errors': [], 'warnings': [], 'entity_info': None}
    ext = os.path.splitext(filepath)[1].lower()

    # PDF入力の場合はバリデーションなしでコピーのみ
    if ext == '.pdf':
        output_name = f'基本契約書_{company_name}.pdf'
        shutil.copy2(filepath, os.path.join(output_dir, output_name))
        result['output_name'] = output_name
        result['warnings'].append('契約書がPDF形式のため、内容チェック・整形処理はスキップされました。')
        return result

    if ext not in ('.docx', '.doc'):
        result['errors'].append(f'契約書: 未対応のファイル形式です ({ext})')
        return result

    output_name = f'基本契約書_{company_name}.docx'
    doc = Document(filepath)
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    # --- 決裁種別チェック ---
    if approval_type == 'paper':
        _check_date_fields(doc, result)
        _check_seal_clause_exists(full_text, result)
    elif approval_type == 'electronic':
        _remove_seal_clause(doc)

    # --- 別紙2の様式チェック ---
    _check_appendix2_version(doc, result)

    # --- カテゴリー及びパートナー ページ削除 ---
    _remove_partner_pages(doc, result)

    # --- 別紙2差し替え ---
    if appendix2_choice:
        _replace_appendix2(doc, appendix2_choice, appendix2_dir, result)

    # --- 書式クリーニング ---
    clean_formatting(doc)

    # --- エンティティ抽出 ---
    result['entity_info'] = extract_entity_info(doc)

    # --- 保存 ---
    output_path = os.path.join(output_dir, output_name)
    doc.save(output_path)
    result['output_name'] = output_name
    return result


def _check_date_fields(doc, result):
    """紙決裁: 年月日欄に具体的な月日が記入されていないかチェック。"""
    for para in doc.paragraphs:
        text = para.text.strip()
        # 「年月日」「年 月 日」パターンを探す
        if re.search(r'年.*月.*日', text):
            # 具体的な月日が入っている場合（例: 2026年4月1日）
            if re.search(r'\d{1,2}\s*月\s*\d{1,2}\s*日', text):
                result['errors'].append(
                    f'【契約書エラー】日付欄に具体的な月日が記入されています: 「{text}」'
                )


def _check_seal_clause_exists(full_text, result):
    """紙決裁: 署名捺印条項の存在確認。"""
    if SEAL_CLAUSE not in full_text:
        # 部分一致でも探す
        if '本契約の成立を証するため' not in full_text:
            result['errors'].append(
                '【契約書エラー】署名捺印条項「本契約の成立を証するため〜」が見つかりません。'
            )


def _remove_seal_clause(doc):
    """電子決裁: 署名捺印条項を削除する。"""
    for para in doc.paragraphs:
        if '本契約の成立を証するため' in para.text:
            para.clear()


def _check_appendix2_version(doc, result):
    """別紙2が最新様式かどうかチェックする。"""
    in_appendix2 = False
    appendix2_text = []

    for para in doc.paragraphs:
        text = para.text.strip()

        # 別紙2セクション開始
        if any(kw in text for kw in APPENDIX2_KEYWORDS):
            in_appendix2 = True
            continue

        # 別紙3に到達したら終了
        if any(kw in text for kw in APPENDIX3_KEYWORDS):
            break

        if in_appendix2:
            appendix2_text.append(text)

    full_appendix2 = '\n'.join(appendix2_text)

    # 旧様式キーワードが含まれていたらエラー
    for old_kw in APPENDIX2_OLD_KEYWORDS:
        if old_kw in full_appendix2:
            result['errors'].append(
                f'【別紙2エラー】旧様式の可能性があります。「{old_kw}」が検出されました。'
                '最新の別紙2に差し替えてください。'
            )
            return

    # 最新様式キーワードが含まれているか確認
    has_latest = any(kw in full_appendix2 for kw in APPENDIX2_LATEST_KEYWORDS)
    if not has_latest and appendix2_text:
        result['warnings'].append(
            '【別紙2確認】最新様式のキーワードが見つかりませんでした。'
            '別紙2が最新版であることを確認してください。'
        )


def _remove_partner_pages(doc, result):
    """カテゴリー及びパートナーのセクションを削除。別紙2・3は保護。"""
    paragraphs_to_remove = []
    in_partner_section = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 別紙2・3に到達したら削除を停止
        if any(kw in text for kw in APPENDIX2_KEYWORDS + APPENDIX3_KEYWORDS):
            in_partner_section = False
            continue

        # パートナーセクション開始検出
        if any(kw in text for kw in PARTNER_PAGE_KEYWORDS):
            in_partner_section = True

        if in_partner_section:
            paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

    if paragraphs_to_remove:
        result['warnings'].append('「カテゴリー及びパートナー」セクションを削除しました。')


def _replace_appendix2(doc, appendix2_filename, appendix2_dir, result):
    """別紙2を指定ファイルに差し替える。"""
    source_path = os.path.join(appendix2_dir, appendix2_filename)
    if not os.path.exists(source_path):
        result['errors'].append(f'別紙2ファイルが見つかりません: {appendix2_filename}')
        return

    try:
        replacement_doc = Document(source_path)
        # 別紙2の開始位置を探す
        start_idx = None
        end_idx = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if any(kw in text for kw in APPENDIX2_KEYWORDS) and start_idx is None:
                start_idx = i
            elif start_idx is not None and any(kw in text for kw in APPENDIX3_KEYWORDS):
                end_idx = i
                break

        if start_idx is not None:
            # 既存の別紙2内容を削除（見出し以降〜別紙3手前まで）
            paras_to_remove = doc.paragraphs[start_idx + 1:end_idx] if end_idx else doc.paragraphs[start_idx + 1:]
            for para in paras_to_remove:
                p = para._element
                p.getparent().remove(p)

            # 差し替え内容を挿入
            insert_after = doc.paragraphs[start_idx]._element
            for para in replacement_doc.paragraphs:
                new_p = para._element.__deepcopy__(True)
                insert_after.addnext(new_p)
                insert_after = new_p

            result['warnings'].append(f'別紙2を「{appendix2_filename}」に差し替えました。')
        else:
            result['warnings'].append('別紙2の位置が特定できませんでした。差し替えをスキップします。')

    except Exception as e:
        result['errors'].append(f'別紙2差し替えエラー: {str(e)}')
