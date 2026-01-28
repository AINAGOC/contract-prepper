"""③ 誓約書の処理"""
import os
import re
from docx import Document

from .common import clean_formatting, extract_entity_info


CORRECT_TITLE = '愛知・名古屋2026大会における大会関係者の宿泊施設等の利用に関する基本契約書'
OLD_TITLE_PATTERNS = [
    r'第20回アジア競技大会.*?基本契約書',
    r'第\d+回アジア競技大会.*?基本契約書',
]


def process_oath(filepath, output_dir, company_name):
    result = {'output_name': '', 'errors': [], 'warnings': [], 'entity_info': None}
    ext = os.path.splitext(filepath)[1].lower()

    # PDF/Excel入力はコピーのみ
    if ext == '.pdf':
        output_name = f'誓約書_{company_name}.pdf'
        import shutil
        shutil.copy2(filepath, os.path.join(output_dir, output_name))
        result['output_name'] = output_name
        result['warnings'].append('誓約書がPDF形式のため、内容チェック・整形処理はスキップされました。')
        return result
    if ext == '.xlsx':
        from openpyxl import load_workbook
        output_name = f'誓約書_{company_name}.xlsx'
        wb = load_workbook(filepath)
        wb.save(os.path.join(output_dir, output_name))
        result['output_name'] = output_name
        return result
    if ext not in ('.docx', '.doc'):
        result['errors'].append(f'誓約書: 未対応のファイル形式です ({ext})')
        return result

    output_name = f'誓約書_{company_name}.docx'
    doc = Document(filepath)

    # --- 件名修正 ---
    _fix_title(doc, result)

    # --- 署名欄チェック ---
    _check_signature(doc, result)

    # --- 書式クリーニング ---
    clean_formatting(doc)

    # --- エンティティ抽出 ---
    result['entity_info'] = extract_entity_info(doc)

    output_path = os.path.join(output_dir, output_name)
    doc.save(output_path)
    result['output_name'] = output_name
    return result


def _fix_title(doc, result):
    """旧件名を新件名に強制置換する。"""
    for para in doc.paragraphs:
        for pattern in OLD_TITLE_PATTERNS:
            if re.search(pattern, para.text):
                old_text = para.text
                for run in para.runs:
                    for pat in OLD_TITLE_PATTERNS:
                        if re.search(pat, run.text):
                            run.text = re.sub(pat, CORRECT_TITLE, run.text)
                result['warnings'].append(
                    f'誓約書の件名を修正しました: 「{old_text.strip()[:30]}...」→ 正式名称'
                )
                return


def _check_signature(doc, result):
    """署名欄が「代表取締役」を含むか確認する。"""
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    if '代表取締役' not in full_text:
        result['errors'].append(
            '【誓約書エラー】署名欄に「代表取締役」の記載がありません。確認してください。'
        )
